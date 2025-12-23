"""Document extraction service with caching and deduplication.

Provides a unified interface for extracting text from PDFs and DOCX files
with automatic caching and cross-parser deduplication.

Usage:
    from components.extraction import DocumentExtractionService
    from components.utils import Cache
    from components.interfaces import Config

    # Extract text from a document
    text = DocumentExtractionService.extract_text(
        url="https://example.com/document.pdf",
        cache=cache_instance,
        config=config_instance
    )

    # Get metrics
    metrics = DocumentExtractionService.get_extraction_metrics()
"""

import io
import logging
import re
import threading
from dataclasses import dataclass
from typing import Optional, TYPE_CHECKING

import PyPDF2  # type: ignore
from docx import Document  # type: ignore

from components.interfaces import _fetch_binary

if TYPE_CHECKING:
    from components.utils import Cache
    from components.interfaces import Config

logger = logging.getLogger(__name__)


@dataclass
class _PendingExtraction:
    """Tracks an in-flight extraction request."""

    event: threading.Event
    result: Optional[str] = None
    error: Optional[Exception] = None


# Thread-safe storage for pending extractions (by URL)
_PENDING_EXTRACTIONS: dict[str, _PendingExtraction] = {}
_PENDING_EXTRACTION_LOCK = threading.RLock()

# Metrics tracking
_EXTRACTION_METRICS = {
    "cache_hits": 0,
    "cache_misses": 0,
    "dedup_waits": 0,
    "extractions": 0,
    "pdf_extractions": 0,
    "docx_extractions": 0,
}
_METRICS_LOCK = threading.Lock()


def get_extraction_metrics() -> dict[str, int]:
    """Get extraction service metrics."""
    with _METRICS_LOCK:
        return _EXTRACTION_METRICS.copy()


def reset_extraction_metrics() -> None:
    """Reset extraction service metrics."""
    with _METRICS_LOCK:
        for key in _EXTRACTION_METRICS:
            _EXTRACTION_METRICS[key] = 0


class DocumentExtractionService:
    """Service for extracting text from documents with caching and
    deduplication.
    """

    @staticmethod
    def extract_text(
        url: str,
        cache: Optional["Cache"] = None,
        config: Optional["Config"] = None,
        timeout: int = 30,
    ) -> Optional[str]:
        """
        Extract text from a document (PDF or DOCX) with caching and
        deduplication.

        This method:
        - Checks cache first (both document cache and extracted text cache)
        - Deduplicates concurrent extraction requests for the same URL
        - Automatically caches extracted text for future use
        - Handles both PDF and DOCX formats

        Args:
            url: URL of the document to extract text from
            cache: Optional cache instance for document and text caching
            config: Optional configuration (required if cache is provided)
            timeout: Timeout for document download in seconds

        Returns:
            Extracted text as string, or None if extraction fails
        """
        # Step 1: Check extracted text cache first (fastest path)
        if cache and config:
            cached_doc = cache.get_cached_document(url, config)
            if cached_doc:
                content_hash = cached_doc.get("content_hash")
                if content_hash:
                    cached_text = cache.get_cached_extracted_text(content_hash, config)
                    if cached_text:
                        with _METRICS_LOCK:
                            _EXTRACTION_METRICS["cache_hits"] += 1
                        logger.debug("Using cached extracted text for %s", url)
                        return cached_text

        # Step 2: Check if extraction is already in progress (deduplication)
        pending_extraction = None
        should_extract = False
        with _PENDING_EXTRACTION_LOCK:
            if url in _PENDING_EXTRACTIONS:
                pending_extraction = _PENDING_EXTRACTIONS[url]
                with _METRICS_LOCK:
                    _EXTRACTION_METRICS["dedup_waits"] += 1
                logger.debug("Waiting for extraction in progress: %s", url)
            else:
                pending_extraction = _PendingExtraction(event=threading.Event())
                _PENDING_EXTRACTIONS[url] = pending_extraction
                should_extract = True

        # Step 3: If extraction is in progress, wait for it
        if not should_extract:
            pending_extraction.event.wait(timeout=timeout + 10)
            if pending_extraction.error:
                logger.warning(
                    "Extraction failed (from other thread): %s - %s",
                    url,
                    pending_extraction.error,
                )
                return None
            if pending_extraction.result:
                return pending_extraction.result
            # If we get here, extraction timed out or failed
            return None

        # Step 4: Perform the extraction
        try:
            with _METRICS_LOCK:
                _EXTRACTION_METRICS["cache_misses"] += 1
                _EXTRACTION_METRICS["extractions"] += 1

            # Determine document type from URL
            is_pdf = bool(re.search(r"\.pdf($|\?)", url, re.I))
            is_docx = bool(re.search(r"\.docx($|\?)", url, re.I))

            if not (is_pdf or is_docx):
                # Try to determine from content type if available
                # For now, default to PDF if uncertain
                logger.debug("Unknown document type for %s, defaulting to PDF", url)
                is_pdf = True

            # Download document (uses document cache if available)
            content = _fetch_binary(url, timeout=timeout, cache=cache, config=config)

            # Extract text based on document type
            if is_pdf:
                extracted_text = DocumentExtractionService._extract_pdf_text(content)
                with _METRICS_LOCK:
                    _EXTRACTION_METRICS["pdf_extractions"] += 1
            elif is_docx:
                extracted_text = DocumentExtractionService._extract_docx_text(content)
                with _METRICS_LOCK:
                    _EXTRACTION_METRICS["docx_extractions"] += 1
            else:
                logger.warning("Unsupported document type for %s", url)
                extracted_text = None

            # Cache the extracted text if we have cache and config
            if extracted_text and cache and config:
                cached_doc = cache.get_cached_document(url, config)
                if cached_doc:
                    content_hash = cached_doc.get("content_hash")
                    if content_hash:
                        cache.cache_extracted_text(content_hash, extracted_text, config)
                        logger.debug("Cached extracted text for %s", url)

            # Store result and notify waiting threads
            with _PENDING_EXTRACTION_LOCK:
                if url in _PENDING_EXTRACTIONS:
                    pending_extraction = _PENDING_EXTRACTIONS.pop(url)
                    pending_extraction.result = extracted_text
                    pending_extraction.event.set()

            return extracted_text

        except Exception as e:  # pylint: disable=broad-exception-caught
            logger.warning("Could not extract text from %s: %s", url, e, exc_info=True)
            # Store error and notify waiting threads
            with _PENDING_EXTRACTION_LOCK:
                if url in _PENDING_EXTRACTIONS:
                    pending_extraction = _PENDING_EXTRACTIONS.pop(url)
                    pending_extraction.error = e
                    pending_extraction.result = None
                    pending_extraction.event.set()
            return None

    @staticmethod
    def _extract_pdf_text(content: bytes) -> Optional[str]:
        """
        Extract text content from PDF bytes.

        Args:
            content: PDF file content as bytes

        Returns:
            Extracted text as string, or None if extraction fails
        """
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_content = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
            if text_content:
                full_text = "\n".join(text_content)
                # Normalize whitespace within lines (preserve newlines)
                # Split by newlines, normalize each line, then rejoin
                lines = full_text.split("\n")
                normalized_lines = [
                    re.sub(r"[ \t]+", " ", line).strip() for line in lines
                ]
                full_text = "\n".join(normalized_lines)
                return full_text
        except Exception as e:  # pylint: disable=broad-exception-caught
            logger.warning("PDF text extraction failed: %s", e)
            return None
        return None

    @staticmethod
    def _extract_docx_text(content: bytes) -> Optional[str]:
        """
        Extract text content from DOCX bytes.

        Extracts text from paragraphs, tables, headers, and footers.

        Args:
            content: DOCX file content as bytes

        Returns:
            Extracted text as string, or None if extraction fails
        """
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            parts = []

            # Extract paragraphs
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text.strip())

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            parts.append(cell.text.strip())

            # Extract headers and footers
            for section in doc.sections:
                if section.header:
                    for p in section.header.paragraphs:
                        if p.text and p.text.strip():
                            parts.append(p.text.strip())
                if section.footer:
                    for p in section.footer.paragraphs:
                        if p.text and p.text.strip():
                            parts.append(p.text.strip())

            docx_file.close()

            if parts:
                # Join parts with spaces (DOCX structure is less line-oriented)
                full_text = " ".join(parts)
                # Normalize whitespace (spaces/tabs, but preserve structure)
                full_text = re.sub(r"[ \t]+", " ", full_text).strip()
                return full_text
        except Exception as e:  # pylint: disable=broad-exception-caught
            logger.warning("DOCX text extraction failed: %s", e)
            return None
        return None
