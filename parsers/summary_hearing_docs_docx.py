"""A parser for DOCX summaries on the hearing's Documents tab."""

import io
import logging
import re
from typing import Optional
from urllib.parse import urljoin, urlparse, parse_qs, unquote

from docx import Document

from components.models import BillAtHearing
from components.interfaces import ParserInterface

logger = logging.getLogger(__name__)

DL_PATH_RX = re.compile(r"/Events/DownloadDocument", re.I)
DOCX_RX = re.compile(r"\.docx($|\?)", re.I)


class SummaryHearingDocsDocxParser(ParserInterface):
    """Parser for DOCX summaries on the hearing's Documents tab."""

    parser_type = ParserInterface.ParserType.SUMMARY
    location = "Hearing page Documents tab DOCX"
    cost = 2

    @staticmethod
    def _norm_bill_id(s: str) -> str:
        """Normalize the bill ID."""
        s = s.upper().replace("\xa0", " ")
        s = re.sub(r"[.\s]", "", s)
        return s

    @staticmethod
    def _title_from_href(href: str) -> str:
        """Get the title from the href."""
        try:
            q = parse_qs(urlparse(href).query)
            t = q.get("Title", [""])[0]
            return unquote(t)
        except Exception:  # pylint: disable=broad-exception-caught
            return ""

    @classmethod
    def _looks_like_summary_for_bill(
        cls, link_text: str, title_param: str, bill_id: str
    ) -> bool:
        """Check if the link looks like a summary for the bill."""
        has_summary = bool(
            re.search(r"\bsummary\b", link_text, re.I) or
            re.search(r"\bsummary\b", title_param, re.I) or
            re.search(r"summary[_\-\s]", link_text, re.I) or
            re.search(r"summary[_\-\s]", title_param, re.I)
        )
        has_bill = (
            bill_id in cls._norm_bill_id(link_text) or
            bill_id in cls._norm_bill_id(title_param)
        )
        return has_summary and has_bill

    @staticmethod
    def _extract_docx_text(docx_url: str) -> Optional[str]:
        """Extract text content from a DOCX URL."""
        try:
            content = ParserInterface._fetch_binary(docx_url, timeout=30)
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            parts = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            parts.append(cell.text.strip())
            for section in doc.sections:
                if section.header:
                    for p in section.header.paragraphs:
                        if p.text and p.text.strip():
                            parts.append(p.text.strip())
                if section.footer:
                    for p in section.footer.paragraphs:
                        if p.text and p.text.strip():
                            parts.append(p.text.strip())
            if parts:
                full_text = " ".join(parts)
                full_text = re.sub(r'\s+', ' ', full_text).strip()
                return full_text
        except Exception as e:  # pylint: disable=broad-exception-caught
            logger.warning("Could not extract text from DOCX %s: %s", docx_url, e)
            return None
        return None

    @classmethod
    def _find_bill_summary_in_docx_text(
        cls, docx_text: str, bill_id: str
    ) -> Optional[str]:
        """Find a bill summary in DOCX text content."""
        lines = docx_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            normalized_line = cls._norm_bill_id(line)
            normalized_bill_id = cls._norm_bill_id(bill_id)
            if (normalized_bill_id in normalized_line and
                    re.search(r'\bsummary\b', line, re.I)):
                return line
        return None

    @classmethod
    def discover(
        cls,
        base_url: str,
        bill: BillAtHearing,
        cache=None,
        config=None
    ) -> Optional[ParserInterface.DiscoveryResult]:
        """
        Probe hearing 'Documents' tab for DOCX summary for this bill.
        Returns {"preview","source_url","confidence"} or None.
        """
        logger.debug("Trying %s...", cls.__name__)
        hearing_docs_url = bill.hearing_url
        soup = cls.soup(hearing_docs_url)
        for a in soup.find_all("a", href=True):
            if not hasattr(a, 'get'):
                continue
            href = a.get("href", "")
            if not isinstance(href, str):
                continue
            if not DL_PATH_RX.search(href):
                continue
            if not DOCX_RX.search(href):
                continue
            docx_url = urljoin(base_url, href)
            text = cls._extract_docx_text(docx_url)
            text = text if text else ""
            title_param = cls._title_from_href(href)
            bill_id = cls._norm_bill_id(bill.bill_id)
            if cls._looks_like_summary_for_bill(
                text, title_param, bill_id
            ):
                preview = (
                    f"Found '{title_param or text}' in hearing "
                    f"Documents for {bill.bill_id}"
                )
                return ParserInterface.DiscoveryResult(
                    preview,
                    "",
                    docx_url,
                    0.95
                )
        for a in soup.find_all("a", href=True):
            if not hasattr(a, 'get'):
                continue
            href = a.get("href", "")
            if not isinstance(href, str):
                continue
            if not DL_PATH_RX.search(href):
                continue
            if not DOCX_RX.search(href):
                continue
            text = " ".join(a.get_text(strip=True).split())
            title_param = cls._title_from_href(href)
            if (re.search(r"\bsummary\b", text, re.I) or
                    re.search(r"\bsummary\b", title_param, re.I) or
                    re.search(r"\breport\b", text, re.I) or
                    re.search(r"\breport\b", title_param, re.I)):
                docx_url = urljoin(base_url, href)
                docx_text = cls._extract_docx_text(docx_url)
                if docx_text:
                    bill_summary_line = (
                        cls._find_bill_summary_in_docx_text(
                            docx_text, bill.bill_id
                        )
                    )
                    if bill_summary_line:
                        preview = (
                            f"Found summary in DOCX content: "
                            f"'{bill_summary_line[:100]}...' "
                            f"for {bill.bill_id}"
                        )
                        return ParserInterface.DiscoveryResult(
                            preview,
                            docx_text,
                            docx_url,
                            0.85,
                        )
        return None

    @staticmethod
    def parse(
        _base_url: str, candidate: ParserInterface.DiscoveryResult
    ) -> dict:
        """Parse the summary."""
        return {"source_url": candidate.source_url}
