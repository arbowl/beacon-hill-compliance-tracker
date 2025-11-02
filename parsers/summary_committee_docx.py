"""A parser for DOCX files in the Committee Summary tab."""

import io
import logging
import re
from typing import Optional
from urllib.parse import urljoin

from bs4 import BeautifulSoup  # type: ignore
from docx import Document

from components.models import BillAtHearing
from components.interfaces import ParserInterface

logger = logging.getLogger(__name__)


class SummaryCommitteeDocxParser(ParserInterface):
    """Parser for DOCX files in the Committee Summary tab."""

    parser_type = ParserInterface.ParserType.SUMMARY
    location = "Committee page Word document"
    cost = 3

    @staticmethod
    def _extract_docx_text(
        docx_url: str,
        cache=None,
        config=None
    ) -> Optional[str]:
        """Extract text content from a DOCX URL (paragraphs, tables, headers,
        footers).
        """
        try:
            content = ParserInterface._fetch_binary(
                docx_url, timeout=30, cache=cache, config=config
            )
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            parts = []
            # Paragraphs
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text.strip())
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            parts.append(cell.text.strip())
            # Headers/Footers
            for section in doc.sections:
                if section.header:
                    for p in section.header.paragraphs:
                        if p.text and p.text.strip():
                            parts.append(p.text.strip())
                if section.footer:
                    for p in section.footer.paragraphs:
                        if p.text and p.text.strip():
                            parts.append(p.text.strip())
            docx_file.close()
            if parts:
                full_text = " ".join(parts)
                full_text = re.sub(r'\s+', ' ', full_text).strip()
                return full_text
        except Exception as e:  # pylint: disable=broad-exception-caught
            logger.warning(
                "Could not extract text from DOCX %s: %s", docx_url, e
            )
            return None
        return None

    @staticmethod
    def _find_committee_summary_docx(
        soup: BeautifulSoup, base_url: str
    ) -> Optional[str]:
        """Find the Committee Summary DOCX link."""
        # Look for any DOCX file on the page - Committee Summary pages
        # typically have only one DOCX file which is the summary
        for a in soup.find_all("a", href=True):
            if not hasattr(a, 'get'):
                continue
            try:
                href = a.get("href", "")
                if not isinstance(href, str):
                    continue
                # Check if it's a DOCX file
                if not re.search(r"\.docx($|\?)", href, re.I):
                    continue
                # Check if it's a download document (typical pattern
                # for summaries)
                if re.search(r"/Download/DownloadDocument/", href, re.I):
                    return urljoin(base_url, href)
                # Also check for committee summary patterns in text or href
                text = a.get_text(strip=True).lower()
                if (
                    re.search(
                        r"committee.*summary|summary.*committee",
                        text,
                        re.I
                    ) or
                    re.search(
                        r"committee.*summary|summary.*committee",
                        href,
                        re.I
                    )
                ):
                    return urljoin(base_url, href)
            except (AttributeError, TypeError):
                continue
        return None

    @classmethod
    def discover(
        cls,
        base_url: str,
        bill: BillAtHearing,
        cache=None,
        config=None
    ) -> Optional[ParserInterface.DiscoveryResult]:
        """Discover the Committee Summary DOCX."""
        logger.debug("Trying %s...", cls.__name__)
        # Navigate to the Committee Summary tab
        committee_summary_url = f"{bill.bill_url}/CommitteeSummary"
        soup = cls.soup(committee_summary_url)
        docx_url = cls._find_committee_summary_docx(soup, base_url)
        if not docx_url:
            return None
        docx_text = None
        if cache and config:
            cached_doc = cache.get_cached_document(docx_url, config)
            if cached_doc:
                content_hash = cached_doc.get("content_hash")
                if content_hash:
                    docx_text = cache.get_cached_extracted_text(
                        content_hash, config
                    )
                    if docx_text:
                        logger.debug(
                            "Using cached extracted text for %s", docx_url
                        )
        if not docx_text:
            docx_text = cls._extract_docx_text(docx_url, cache, config)
            if docx_text and cache and config:
                cached_doc = cache.get_cached_document(docx_url, config)
                if cached_doc:
                    content_hash = cached_doc.get("content_hash")
                    if content_hash:
                        cache.cache_extracted_text(
                            content_hash, docx_text, config
                        )
                        logger.debug("Cached extracted text for %s", docx_url)
        if docx_text:
            # Use the extracted text as preview, truncated if too long
            preview = f"Found Committee Summary DOCX for {bill.bill_id}"
            if len(docx_text) > 200:
                preview += f"\n\nDOCX Content Preview:\n{docx_text[:500]}..."
            else:
                preview += f"\n\nDOCX Content:\n{docx_text}"
            return ParserInterface.DiscoveryResult(
                preview,
                docx_text,
                docx_url,
                0.9,
            )
        else:
            # Fallback to simple preview if text extraction fails
            preview = (
                f"Found Committee Summary DOCX for {bill.bill_id} "
                "(text extraction failed)"
            )
            return ParserInterface.DiscoveryResult(
                preview,
                "",
                docx_url,
                0.8,  # Lower confidence if we can't extract text
            )

    @staticmethod
    def parse(
        _base_url: str, candidate: ParserInterface.DiscoveryResult
    ) -> dict:
        """Parse the Committee Summary DOCX."""
        # For now, just return the URL we confirmed
        # Later we can add actual DOCX text extraction if needed
        return {"source_url": candidate.source_url}
